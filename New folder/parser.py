"""
Parser Module for AI Resume Screening Agent.
Handles robust text extraction from PDF, DOCX, and TXT files.
Includes fallbacks, encoding detection, and defensive error handling.
"""

import os
import io
import logging
from typing import Union, BinaryIO

import pdfplumber
import PyPDF2
import docx

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)


class DocumentParser:
    """Document Parser supporting PDF, DOCX, and TXT formats with fallback mechanisms."""

    @staticmethod
    def parse_pdf(file_source: Union[str, BinaryIO, bytes]) -> str:
        """
        Extract text from PDF using pdfplumber with PyPDF2 fallback.
        Handles corrupt PDFs, scans without text, and encoding anomalies.
        """
        text = ""

        # Handle bytes vs filename vs file object
        if isinstance(file_source, bytes):
            stream = io.BytesIO(file_source)
        elif isinstance(file_source, str) and os.path.exists(file_source):
            with open(file_source, "rb") as f:
                stream = io.BytesIO(f.read())
        elif hasattr(file_source, "read"):
            content = file_source.read()
            if hasattr(file_source, "seek"):
                file_source.seek(0)
            stream = io.BytesIO(content)
        else:
            logger.error("Invalid file source provided to parse_pdf")
            return ""

        # Primary attempt: pdfplumber
        try:
            stream.seek(0)
            with pdfplumber.open(stream) as pdf:
                pages_text = []
                for page_idx, page in enumerate(pdf.pages):
                    extracted = page.extract_text()
                    if extracted:
                        pages_text.append(extracted)
                text = "\n".join(pages_text).strip()
                if text:
                    logger.info(f"Successfully extracted {len(text)} chars using pdfplumber.")
                    return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}. Attempting PyPDF2 fallback...")

        # Fallback attempt: PyPDF2
        try:
            stream.seek(0)
            reader = PyPDF2.PdfReader(stream)
            pages_text = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted)
            text = "\n".join(pages_text).strip()
            if text:
                logger.info(f"Successfully extracted {len(text)} chars using PyPDF2 fallback.")
                return text
        except Exception as e:
            logger.error(f"PyPDF2 fallback also failed: {e}")

        return text

    @staticmethod
    def parse_docx(file_source: Union[str, BinaryIO, bytes]) -> str:
        """Extract text from Microsoft Word (.docx) files."""
        try:
            if isinstance(file_source, bytes):
                stream = io.BytesIO(file_source)
            elif isinstance(file_source, str) and os.path.exists(file_source):
                stream = file_source
            elif hasattr(file_source, "read"):
                content = file_source.read()
                if hasattr(file_source, "seek"):
                    file_source.seek(0)
                stream = io.BytesIO(content)
            else:
                return ""

            doc = docx.Document(stream)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            text = "\n".join(full_text).strip()
            logger.info(f"Successfully extracted {len(text)} chars from DOCX.")
            return text
        except Exception as e:
            logger.error(f"Failed to parse DOCX file: {e}")
            return ""

    @staticmethod
    def parse_txt(file_source: Union[str, BinaryIO, bytes]) -> str:
        """Extract text from plain text (.txt) files with multiple encoding fallbacks."""
        try:
            if isinstance(file_source, bytes):
                raw_bytes = file_source
            elif isinstance(file_source, str) and os.path.exists(file_source):
                with open(file_source, "rb") as f:
                    raw_bytes = f.read()
            elif hasattr(file_source, "read"):
                raw_bytes = file_source.read()
                if hasattr(file_source, "seek"):
                    file_source.seek(0)
            else:
                return ""

            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
            for enc in encodings:
                try:
                    text = raw_bytes.decode(enc).strip()
                    logger.info(f"Parsed TXT file using {enc} encoding.")
                    return text
                except UnicodeDecodeError:
                    continue

            return raw_bytes.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            logger.error(f"Failed to parse TXT file: {e}")
            return ""

    @classmethod
    def parse_document(cls, file_source: Union[str, BinaryIO, bytes], filename: str = "") -> str:
        """
        Unified document parser entry point. Automatically detects format from extension or content.
        """
        filename_lower = filename.lower()

        if filename_lower.endswith(".pdf"):
            return cls.parse_pdf(file_source)
        elif filename_lower.endswith(".docx"):
            return cls.parse_docx(file_source)
        elif filename_lower.endswith(".txt"):
            return cls.parse_txt(file_source)
        else:
            # Attempt to infer or try PDF then DOCX then TXT
            logger.warning(f"Unknown extension for '{filename}'. Attempting auto-detection...")
            text = cls.parse_pdf(file_source)
            if not text:
                text = cls.parse_docx(file_source)
            if not text:
                text = cls.parse_txt(file_source)
            return text
